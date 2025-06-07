from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import re

from .utils import *
from .style import cs_lab_report

class DocxHandlerGenerator:
    """
    WHU CS 实验报告 - docx 生成器类。
    """
    dpi = 96
    def __init__(self):
        self.max_image_width_inches = None
        self.document = Document()
        # init
        self.init()
        # Style
        cs_lab_report.set_style(self.document)

    def init(self):
        section = self.document.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin  # TIP 单位：EMU
        # 计算图片的最大宽度
        max_image_width_emu = page_width - left_margin - right_margin
        self.max_image_width_inches = max_image_width_emu / 914400  # EMU to inches conversion (1 inch = 914400 EMU)

    def _is_last_element_title(self):
        """
        检查文档中的最后一个元素是否为标题
        :return: 如果最后一个元素是标题则返回True，否则返回False
        """
        if len(self.document.paragraphs) == 0:
            return False
        
        last_paragraph = self.document.paragraphs[-1]
        # 检查段落的样式是否包含'title'关键字
        if last_paragraph.style.name and 'title' in last_paragraph.style.name.lower():
            return True
        return False

    def add_title(self, idx, title, level):
        """
        添加 CS 要求的标题格式；
        :param idx: 由 MD-AST 计算得出。
        :param title:
        :param level:
        :return:
        """
        if level == 1:
            self.add_page_break()

        # 二级标题换页
        if level == 2 and (idx[1] > 1 or (idx[1] == 1 and not self._is_last_element_title())):
            self.document.add_paragraph()

        # CORE
        pa = self.document.add_paragraph(".".join(str(i) for i in idx), style=f'title-{level}-num')
        # ADD：一个前置空格
        pa.add_run(cs_lab_report.TITLE_MID + title, style=f'title-{level}').bold = True

        # UPDATE 一级标题居中
        if level == 1:
            pa.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_text(self, text):
        self.document.add_paragraph(cs_lab_report.TEXT_PREFIX + text, style='body')

    def add_text_with_strikethrough(self, text):
        """
        处理包含删除线标记的文本
        :param text: 可能包含~~标记的文本
        """
        p = self.document.add_paragraph(style='body')
        p.add_run(cs_lab_report.TEXT_PREFIX)
        
        # 查找所有的删除线部分
        parts = []
        start = 0
        while '~~' in text[start:]:
            # 找到第一个 ~~
            first_pos = text.find('~~', start)
            if first_pos > start:
                # 添加前面的普通文本
                parts.append((text[start:first_pos], False))
            
            # 找到第二个 ~~
            second_pos = text.find('~~', first_pos + 2)
            if second_pos != -1:
                # 添加删除线文本
                strikethrough_text = text[first_pos + 2:second_pos]
                parts.append((strikethrough_text, True))
                start = second_pos + 2
            else:
                # 没有找到第二个~~，将剩余部分作为普通文本
                parts.append((text[first_pos:], False))
                break
        
        # 添加最后的普通文本
        if start < len(text):
            parts.append((text[start:], False))
        
        # 添加各部分到段落
        for part_text, is_strikethrough in parts:
            run = p.add_run(part_text)
            if is_strikethrough:
                run.font.strike = True

    def add_paragraph_with_runs(self, children):
        """
        添加包含多种格式的段落
        :param children: 段落中的子元素列表
        """
        p = self.document.add_paragraph(style='body')
        
        # 添加段落前缀
        p.add_run(cs_lab_report.TEXT_PREFIX)
        
        for child in children:
            # 打印类型信息，用于调试
            print(f"处理段落中的元素类型: {child['type']}")
            if 'raw' in child:
                print(f"内容: {child['raw']}")
            if 'children' in child and child['children']:
                first_child = child['children'][0]
                if 'raw' in first_child:
                    print(f"子元素内容: {first_child['raw']}")
                    
            if child['type'] == 'text':
                # 检查文本是否包含删除线标记
                text = child['raw']
                if '~~' in text:
                    # 使用更复杂的处理逻辑来处理可能包含多个删除线标记的文本
                    start = 0
                    while '~~' in text[start:]:
                        # 找到第一个 ~~
                        first_pos = text.find('~~', start)
                        if first_pos > start:
                            # 添加前面的普通文本
                            p.add_run(text[start:first_pos])
                        
                        # 找到第二个 ~~
                        second_pos = text.find('~~', first_pos + 2)
                        if second_pos != -1:
                            # 添加删除线文本
                            strikethrough_text = text[first_pos + 2:second_pos]
                            run = p.add_run(strikethrough_text)
                            run.font.strike = True
                            start = second_pos + 2
                        else:
                            # 没有找到第二个~~，将剩余部分作为普通文本
                            p.add_run(text[start:])
                            break
                    
                    # 添加最后的普通文本
                    if start < len(text):
                        p.add_run(text[start:])
                else:
                    p.add_run(text)
            elif child['type'] == 'strong':
                if 'children' in child and child['children']:
                    run = p.add_run(child['children'][0]['raw'])
                    run.bold = True
            elif child['type'] == 'emphasis':
                if 'children' in child and child['children']:
                    child_elem = child['children'][0]
                    if child_elem['type'] == 'text':
                        run = p.add_run(child_elem['raw'])
                        run.italic = True
                    elif child_elem['type'] == 'strong':
                        # 处理粗斜体
                        if 'children' in child_elem and child_elem['children']:
                            run = p.add_run(child_elem['children'][0]['raw'])
                            run.bold = True
                            run.italic = True
            elif child['type'] == 'codespan':
                if 'raw' in child:
                    run = p.add_run(child['raw'])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
            elif child['type'] == 's' or child['type'] == 'strikethrough' or child['type'] == 'del':
                # 处理各种可能的删除线表示
                if 'children' in child and child['children']:
                    run = p.add_run(child['children'][0]['raw'])
                    run.font.strike = True
                elif 'raw' in child:
                    run = p.add_run(child['raw'])
                    run.font.strike = True
            elif child['type'] == 'softbreak':
                p.add_run('\n')
            elif child['type'] == 'link':
                if 'children' in child and child['children'] and 'attrs' in child:
                    text = child['children'][0].get('raw', '')
                    url = child['attrs'].get('url', '')
                    
                    # 为超链接创建单独的run
                    run = p.add_run(text)
                    run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                    run.font.underline = True
                    
                    # 添加超链接关系
                    r_id = f'rId{len(self.document.part._rels) + 1}'
                    self.document.part._rels.add_relationship(
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                        url,
                        is_external=True,
                        rId=r_id
                    )
                    
                    # 设置超链接属性
                    for e in run._element.xpath('./w:rPr'):
                        run._element.remove(e)
                        
                    # 创建全新的超链接元素
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    
                    # 创建运行元素
                    new_run = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    
                    # 设置样式
                    rStyle = OxmlElement('w:rStyle')
                    rStyle.set(qn('w:val'), 'Hyperlink')
                    rPr.append(rStyle)
                    
                    # 设置颜色
                    color = OxmlElement('w:color')
                    color.set(qn('w:val'), '0000FF')
                    rPr.append(color)
                    
                    # 设置下划线
                    u = OxmlElement('w:u')
                    u.set(qn('w:val'), 'single')
                    rPr.append(u)
                    
                    new_run.append(rPr)
                    
                    # 添加文本
                    t = OxmlElement('w:t')
                    t.text = text
                    new_run.append(t)
                    
                    # 组装超链接
                    hyperlink.append(new_run)
                    
                    # 替换原来的run
                    p._p.append(hyperlink)

    def add_bold_text(self, text):
        p = self.document.add_paragraph(style='body')
        p.add_run(text).bold = True

    def add_italic_text(self, text):
        p = self.document.add_paragraph(style='body')
        p.add_run(text).italic = True

    def add_strikethrough_text(self, text):
        p = self.document.add_paragraph(style='body')
        run = p.add_run(text)
        run.font.strike = True

    def add_inline_code(self, text):
        p = self.document.add_paragraph(style='body')
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_code_block(self, code, language=''):
        # 创建一个有边框和背景色的表格来包含代码
        table = self.document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        # 设置表格属性
        table.autofit = False
        table.allow_autofit = False
        
        # 设置单元格属性
        cell = table.cell(0, 0)
        cell.width = self.document.sections[0].page_width - self.document.sections[0].left_margin - self.document.sections[0].right_margin - Pt(40)
        
        # 设置背景色为浅灰色
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # 添加代码文本
        paragraph = cell.paragraphs[0]
        paragraph.style = 'Normal'
        run = paragraph.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)  # 稍微小一点的字体
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 减小段落间距
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(5)
        
        # 如果有语言标签，添加到代码块上方
        if language and language.strip():
            lang_para = self.document.add_paragraph(style='body')
            lang_para.paragraph_format.space_after = Pt(0)
            lang_run = lang_para.add_run(language.strip())
            lang_run.font.size = Pt(8)
            lang_run.font.color.rgb = RGBColor(128, 128, 128)
            lang_run.font.italic = True
            
        # 添加表格周围的空间
        self.document.add_paragraph().paragraph_format.space_after = Pt(10)

    def add_quote(self, children):
        p = self.document.add_paragraph(style='body')
        p.paragraph_format.left_indent = Pt(20)
        p.paragraph_format.right_indent = Pt(20)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)
        p.add_run('"')
        for child in children:
            if child['type'] == 'paragraph':
                for subchild in child['children']:
                    if subchild['type'] == 'text':
                        p.add_run(subchild['raw'])
                    elif subchild['type'] == 'strong':
                        p.add_run(subchild['children'][0]['raw']).bold = True
                    elif subchild['type'] == 'emphasis':
                        p.add_run(subchild['children'][0]['raw']).italic = True
        p.add_run('"')

    def add_bullet_list(self, items):
        for item in items:
            p = self.document.add_paragraph(style='body')
            p.paragraph_format.left_indent = Pt(20)
            p.add_run('• ').bold = True
            for child in item['children']:
                if child['type'] == 'block_text':
                    for subchild in child['children']:
                        if subchild['type'] == 'text':
                            p.add_run(subchild['raw'])
                        elif subchild['type'] == 'strong':
                            p.add_run(subchild['children'][0]['raw']).bold = True
                        elif subchild['type'] == 'emphasis':
                            p.add_run(subchild['children'][0]['raw']).italic = True

    def add_ordered_list(self, items):
        for i, item in enumerate(items, 1):
            p = self.document.add_paragraph(style='body')
            p.paragraph_format.left_indent = Pt(20)
            p.add_run(f'{i}. ').bold = True
            for child in item['children']:
                if child['type'] == 'block_text':
                    for subchild in child['children']:
                        if subchild['type'] == 'text':
                            p.add_run(subchild['raw'])
                        elif subchild['type'] == 'strong':
                            p.add_run(subchild['children'][0]['raw']).bold = True
                        elif subchild['type'] == 'emphasis':
                            p.add_run(subchild['children'][0]['raw']).italic = True

    def add_table(self, children):
        """
        添加表格 - 当前版本仅显示提示信息
        :param children: 表格数据
        """
        # 输出警告信息
        print(f"⚠️ 警告: 当前版本不支持完整的表格渲染")
        
        # 添加提示段落
        p = self.document.add_paragraph(style='body')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表格内容请在Markdown文件中查看")
        run.bold = True
        run.font.size = Pt(10)
        
        # 添加空间
        self.document.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_math(self, formula, is_inline=False):
        """
        添加数学公式，当前版本仅以文本形式输出
        :param formula: LaTeX格式的公式
        :param is_inline: 是否为行内公式
        """
        # 输出警告信息
        print(f"⚠️ 警告: 当前版本不支持完整的LaTeX公式渲染: {formula}")
        
        # 清理LaTeX公式，去除$和$$符号
        formula = formula.strip()
        if formula.startswith('$$') and formula.endswith('$$'):
            formula = formula[2:-2].strip()
        elif formula.startswith('$') and formula.endswith('$'):
            formula = formula[1:-1].strip()
        
        # 简单呈现公式
        if not is_inline:
            # 块级公式居中显示
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 简单的表格容器
            table = self.document.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            
            # 设置单元格
            cell = table.cell(0, 0)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加公式文本
            run = para.add_run(formula)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
            run.font.italic = True
            
            # 添加空间
            self.document.add_paragraph().paragraph_format.space_after = Pt(6)
        else:
            # 行内公式
            p = self.document.add_paragraph(style='body')
            p.add_run(" ")  # 添加前导空格
            
            # 添加行内公式
            math_run = p.add_run(formula)
            math_run.font.italic = True
            
            p.add_run(" ")  # 添加尾随空格

    def add_hyperlink(self, text, url):
        """
        添加超链接
        :param text: 显示的文本
        :param url: 链接地址
        """
        print(f"⚠️ 警告: 超链接功能未实现: {text} -> {url}")
        return
        p = self.document.add_paragraph(style='body')
        
        # 创建超链接关系
        hyperlink = OxmlElement('w:hyperlink')
        
        # 创建关系ID
        r_id = f'rId{len(self.document.part._rels) + 1}'
        hyperlink.set(qn('r:id'), r_id)
        
        # 添加关系
        self.document.part._rels.add_relationship(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            url,
            is_external=True,
            rId=r_id
        )
        
        # 创建运行元素
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # 设置样式
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        
        # 设置颜色
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        # 设置下划线
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        # 添加文本
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        # 组装超链接
        hyperlink.append(new_run)
        
        # 添加到段落
        p._p.append(hyperlink)

    def add_picture_with_text(self, idx, img_url, text):
        pa = self.document.add_paragraph()
        run = pa.add_run()
        # 图
        image_data, image_width, image_height = download_image(img_url)
        run.add_picture(image_data, width=calculate_image_width(image_width, self.max_image_width_inches, self.dpi))
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 图例文本
        pa = self.document.add_paragraph(cs_lab_report.IMG_TEXT_PREFIX, style='img-title')
        pa.add_run(".".join(str(i) for i in idx), style='img-title-num').bold = True
        pa.add_run(cs_lab_report.IMG_TEXT_MID + text, style='img-title-2')
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_page_break(self):
        self.document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def save(self, filename):
        self.document.save(filename)


if __name__ == '__main__':
    docx_handler = DocxHandlerGenerator()
    docx_handler.add_title([2], '这是一个标题', 1)
    docx_handler.add_title([1, 1], "H2", 2)
    docx_handler.add_text(
        "测试文本段覅哦沙发第四哦分段覅哦沙发第四哦分段覅哦沙发第四哦分，段覅哦沙发第四哦分段覅哦沙发第四哦分。")

    docx_handler.add_picture_with_text([1, 1], "https://www.helloimg.com/i/2024/12/25/676b92d11041e.png", "测试图片")

    docx_handler.save('example.docx')
