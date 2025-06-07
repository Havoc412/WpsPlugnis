import cv2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from style import cs_lab_report
from utils import *

class DocxHandler:
    dpi = 96
    def __init__(self):
        self.document = Document()
        self.max_image_width_inches = None
        # init
        self.init()
        # Style
        cs_lab_report.set_style(self.document)

    def init(self):
        # 获取文档的页面宽度
        section = self.document.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin  # TIP 单位：EMU

        # 计算图片的最大宽度
        max_image_width_emu = page_width - left_margin - right_margin
        self.max_image_width_inches = max_image_width_emu / 914400  # EMU to inches conversion (1 inch = 914400 EMU)

    def add_custom_title(self, idx, title, level):
        # # TIP 创建自定义样式; dict 的方式，更适合集中处理出来。
        # style = self.document.styles.add_style(f'Title {level}', 1)  # 1 表示段落样式
        # style.font.name = '黑体'
        # style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # style.font.size = Pt(22)  # 2号字大约是22磅

        # 添加段落并应用样式
        pa = self.document.add_paragraph(str(idx), style=f'title-{level}-num')
        # ADD：一个前置空格
        pa.add_run(cs_lab_report.TITLE_MID + title, style=f'title-{level}')

        if level == 1:
            pa.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_text(self, text):
        self.document.add_paragraph(cs_lab_report.TEXT_PREFIX + text, style='body')

    def add_picture_with_text(self, img_url, idx, text):
        pa = self.document.add_paragraph()
        run = pa.add_run()
        # 图
        image_data, image_width, image_height = download_image(img_url)
        run.add_picture(image_data, width=calculate_image_width(image_width, self.max_image_width_inches, self.dpi))
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 图例文本
        pa = self.document.add_paragraph(cs_lab_report.IMG_TEXT_PREFIX, style='img-title')
        pa.add_run(str(idx), style='img-title-num')
        pa.add_run(cs_lab_report.IMG_TEXT_MID + text, style='img-title-2')
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save(self, filename):
        self.document.save(filename)

# 示例用法
docx_handler = DocxHandler()
docx_handler.add_custom_title(2, '这是一个标题', 1)
docx_handler.add_custom_title(2.1, "H2", 2)
docx_handler.add_text("测试文本段覅哦沙发第四哦分段覅哦沙发第四哦分段覅哦沙发第四哦分，段覅哦沙发第四哦分段覅哦沙发第四哦分。")

docx_handler.add_picture_with_text("https://www.helloimg.com/i/2024/12/21/6765ed167c8f8.png", 1.1, "测试图片")

docx_handler.save('example.docx')
