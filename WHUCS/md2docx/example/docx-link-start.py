from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import cv2
import os

# 创建一个新的Word文档对象
document = Document()

# 获取文档的页面宽度
section = document.sections[0]
page_width = section.page_width
left_margin = section.left_margin
right_margin = section.right_margin  # TIP 单位：EMU

# 计算图片的最大宽度
max_image_width_emu = page_width - left_margin - right_margin
max_image_width_inches = max_image_width_emu / 914400  # EMU to inches conversion (1 inch = 914400 EMU)


# 插入一级标题 "Document Title"
document.add_heading('Document Title', level=0)  # TIP 实测在 WPS-docx 中不属于 Hx 系类。

# TAG 插入一个普通段落，并在其中添加加粗和斜体格式的文本
paragraph = document.add_paragraph('A plain paragraph having some ')
paragraph.add_run('bold').bold = True  # 加粗文本 "bold"
paragraph.add_run(' and some ')
paragraph.add_run('italic.').italic = True  # 斜体文本 "italic."

# TAG 插入一级标题 "Heading, level 1"
document.add_heading('Heading, level 1', level=1)

# 插入带有特殊样式的引用段落 "Intense quote"
document.add_paragraph('Intense quote', style='Intense Quote')

# 插入 TAG 无序列表 的第一项 "first item in unordered list"
document.add_paragraph('first item in unordered list', style='List Bullet')

# 插入 TAG 有序列表 的第一项 "first item in ordered list"
document.add_paragraph('first item in ordered list', style='List Number')

# 图片路径
image_path = './test.jpeg'  # 小图片
# image_path = "./中文.jpeg"  # BUG opencv2 读取中文路径之时，需要 np 的预处理。
# image_path = r"C:\Users\Havoc\Pictures\Screenshots\中文.png"  # 大图片
if not os.path.exists(image_path):
    print("Image not found.")
    exit(1)

# 使用 OpenCV 打开图片以获取其原始宽度
img = cv2.imread(image_path)
original_height, original_width, _ = img.shape

dpi = 96
original_width_inches = original_width / dpi

# 计算图片的实际宽度（英寸）
if original_width_inches > max_image_width_inches:
    image_width_inches = max_image_width_inches
else:
    image_width_inches = original_width_inches

# 将英寸转换为 EMU，以便插入图片时使用
image_width_emu = Inches(image_width_inches).emu

# 添加图片到段落，并设置宽度
paragraph = document.add_paragraph()
run = paragraph.add_run()
run.add_picture(image_path, width=image_width_emu)

# 设置段落对齐方式为居中
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

paragraph = document.add_paragraph("图1.2 Test 图例")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 定义表格数据
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# 插入一个包含三列的表格，并填充表头
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'  # 表头第一列 "Qty"
hdr_cells[1].text = 'Id'   # 表头第二列 "Id"
hdr_cells[2].text = 'Desc' # 表头第三列 "Desc"

# 循环插入表格数据行
for qty, id, desc in records:
    row_cells = table.add_row().cells  # 新增一行
    row_cells[0].text = str(qty)       # 填充 "Qty" 列
    row_cells[1].text = id             # 填充 "Id" 列
    row_cells[2].text = desc           # 填充 "Desc" 列

# 插入分页符
document.add_page_break()

# 将文档保存为 "test.docx" 文件
document.save('test.docx')
